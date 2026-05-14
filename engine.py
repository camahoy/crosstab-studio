"""
engine.py — Crosstab Studio v1.1
Reliable reading of Corporate Reputation, Global Brand Identity, and KP formats.
Wave-by-wave comparison support.
"""

print("CROSSTAB STUDIO ENGINE v1.1")

import io, math, re, os, json
from collections import Counter
import numpy as np
import pandas as pd
import openpyxl
from openpyxl.styles import Font as XLFont, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from profiles import PROFILES, get_profile as _get_profile

def _profile(name):
    """Resolve a profile by name, checking user profiles too."""
    p = _get_profile(name)
    if p is None:
        raise KeyError(f"Unknown profile: {name!r}. Check user_profiles.json.")
    return p

USER_PROFILES_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'user_profiles.json')


def load_user_profiles():
    try:
        if os.path.exists(USER_PROFILES_PATH):
            with open(USER_PROFILES_PATH, 'r') as f:
                return json.load(f)
    except Exception:
        pass
    return {}


def save_user_profile(name, profile_dict):
    profiles = load_user_profiles()
    profiles[name] = profile_dict
    with open(USER_PROFILES_PATH, 'w') as f:
        json.dump(profiles, f, indent=2)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph as DocxPara


# ── Value coercion ────────────────────────────────────────────

def _coerce(v):
    """Convert any cell value to float if possible, None if blank, string otherwise."""
    if v is None: return None
    if isinstance(v, float): return None if math.isnan(v) else v
    if isinstance(v, int):   return float(v)
    if isinstance(v, str):
        s = v.strip()
        if s in ('', '-', '\xa0', '\u00a0'): return None
        try:    return float(s)
        except: return v
    return v


# ── Format validator ──────────────────────────────────────────

def validate_format(file_bytes):
    """
    Check uploaded file against all known profiles.
    Returns (matched_profile_name, confidence) or (None, 0).
    """
    try:
        xl  = pd.ExcelFile(io.BytesIO(file_bytes))
        si  = 1 if len(xl.sheet_names) > 1 else 0
        df  = xl.parse(si, header=None, nrows=12, na_values=[''])
        raw = df.values.tolist()
    except Exception:
        return None, 0

    def cell(r, c):
        try:
            v = raw[r][c]
            return str(v).strip() if v and str(v) not in ('nan','None') else ''
        except: return ''

    def row_has_total(r):
        """Check if any cell in a row contains total."""
        try:
            return any('total' in str(v).lower()
                       for v in (raw[r] or [])
                       if v and str(v) not in ('nan','None'))
        except: return False

    # Corporate Reputation (fmt6):
    # Row 2 = descriptor ('Total sample'), row 3 = question, row 4 col 1 = Total
    if len(raw) > 5:
        r2   = cell(2, 0)
        r3   = cell(3, 0)
        r4c1 = cell(4, 1)
        r5c1 = cell(5, 1)
        if (len(r3) > 10
                and ('total sample' in r2.lower() or 'weight' in r2.lower() or len(r2) < 50)
                and ('total' in r4c1.lower() or 'total' in r5c1.lower())):
            return "Corporate Reputation", 95

    # Global Brand Identity (fmt2):
    # Row 2 = question, row 3 col 0 or col 1 = Total, row 4 = sub-labels
    if len(raw) > 5:
        r2   = cell(2, 0)
        r3c0 = cell(3, 0)
        r3c1 = cell(3, 1)
        r4c1 = cell(4, 1)
        if (len(r2) > 10
                and ('total' in r3c0.lower() or 'total' in r3c1.lower())
                and 'total' in r4c1.lower()):
            return "Global Brand Identity", 90

    # KP Omni: row 1=project name, row 2="Table N", row 4=question
    # Has "Base Weighted" somewhere in rows 10-20
    if len(raw) > 6:
        r2 = cell(2, 0)
        r4 = cell(4, 0)
        if (r2.lower().startswith('table') and len(r4) > 10):
            # Check for Base Weighted
            has_base_weighted = any(
                'base weighted' in str(raw[ri][0] if raw[ri] else '').lower()
                for ri in range(min(20, len(raw)))
            )
            if has_base_weighted:
                return "KP Omni", 92

    # KP: row 2 = question, Total somewhere in row 1
    if len(raw) > 3:
        r2 = cell(2, 0)
        if len(r2) > 10 and row_has_total(1):
            r4c1 = cell(4, 1)
            if 'total' not in r4c1.lower():
                return "KP", 85

    return None, 0


# ── Fast scanner ──────────────────────────────────────────────

def fast_scan(file_bytes, profile_name):
    """
    Rapidly read question names from each sheet.
    Returns ordered list of question groups.
    """
    profile = _profile(profile_name)
    q_row   = profile["question_row"]
    xl      = pd.ExcelFile(io.BytesIO(file_bytes))
    groups  = {}
    order   = []

    for i, sname in enumerate(xl.sheet_names):
        # Skip sheet 0 if it looks like a TOC/index
        if i == 0 and profile.get("skip_sheet_0"):
            try:
                df0  = xl.parse(0, header=None, nrows=max(q_row + 2, 5), na_values=[''])
                raw0 = df0.values.tolist()
                # Detect TOC by known header strings in first few rows
                toc_indicators = ['sheet number', 'table of contents', 'contents', 'index']
                is_toc = any(
                    isinstance(raw0[r][0], str) and
                    any(ind in raw0[r][0].strip().lower() for ind in toc_indicators)
                    for r in range(min(4, len(raw0)))
                    if raw0[r] and raw0[r][0]
                )
                c0 = str(raw0[q_row][0]).strip() if len(raw0) > q_row and raw0[q_row] and raw0[q_row][0] else ''
                if is_toc or not c0 or len(c0) < 5:
                    continue
            except Exception:
                continue

        try:
            df  = xl.parse(i, header=None, nrows=q_row + 8, na_values=[''])
            raw = df.values.tolist()
            if len(raw) <= q_row: continue

            # Get question wording — KP may span multiple rows
            wording = str(raw[q_row][0]).strip() if raw[q_row] and raw[q_row][0] else ''
            if not wording or len(wording) < 5: continue
            if any(x in wording.lower() for x in ['sample', 'weight', 'project id', 'table of', 'table 1', 'table 2', 'table 3', 'table 4', 'table 5', 'table 6']): continue

            # For KP Omni: q_row=4 which is the question, but row 2 is "Table N"
            # We want to use the question text not the table number
            # Also grab statement from row 5 to append
            if profile_name == "KP Omni":
                stmt = _get_kp_omni_statement(raw)
                if stmt:
                    wording = wording + " — " + stmt

            # For KP: check if wording continues on next rows
            if profile_name == "KP":
                for extra_row in range(q_row + 1, min(q_row + 7, len(raw))):
                    next_cell = str(raw[extra_row][0]).strip() if raw[extra_row] and raw[extra_row][0] else ''
                    if not next_cell or 'base' in next_cell.lower() or '=' in next_cell:
                        break
                    wording += ' ' + next_cell

            # Detect sheet type from wording
            sheet_type = _detect_sheet_type(wording)

            # Extract prefix
            m      = re.match(r'^([A-Za-z0-9_]+)[\.\s]', wording)
            prefix = m.group(1) if m else wording[:8]

            if prefix not in groups:
                groups[prefix] = {
                    'prefix':     prefix,
                    'wording':    wording[:100],
                    'sheets':     [],
                    'types':      set(),
                }
                order.append(prefix)

            groups[prefix]['sheets'].append(i)
            groups[prefix]['types'].add(sheet_type)

        except Exception:
            continue

    # Convert types set to sorted list
    for g in groups.values():
        g['types'] = sorted(g['types'])

    return [groups[p] for p in order]


def _detect_sheet_type(wording):
    """Classify a sheet as standard, t2b, b2b, summary_grid, or mean."""
    wl = wording.lower().replace("'[","[")
    if 't2b' in wl or "top 2 box" in wl:         return 't2b'
    if 'b2b' in wl or "bottom 2 box" in wl:       return 'b2b'
    if 'summary grid' in wl or 'grid' in wl:       return 'summary_grid'
    if 'mean' in wl or 'average' in wl:            return 'mean'
    return 'standard'


# ── Column reader ─────────────────────────────────────────────

def _find_header_row(raw, max_rows=35):
    """
    Find the row that contains 'Total' as a column header value.
    Skips rows where the majority of non-blank values are single letters (sig-code rows).
    Returns row index or None.
    """
    for ri, row in enumerate(raw[:max_rows]):
        if not row:
            continue
        str_vals = [str(v).strip() for v in row
                    if v is not None and str(v).strip() not in ('', 'nan', 'None', '\xa0')]
        if not str_vals:
            continue
        # Skip letter-code rows (>= half the values are single alpha chars)
        single_letter = sum(1 for s in str_vals if len(s) == 1 and s.isalpha())
        if single_letter >= len(str_vals) / 2:
            continue
        # Look for an exact 'total' match in any column
        for v in str_vals:
            if v.lower() == 'total' or v.lower().startswith('total respondent'):
                return ri
    return None


def get_columns(file_bytes, profile_name):
    """
    Get available column names from the first real data sheet.
    Finds the header row by locating the row that contains 'Total' as a
    column header value, skipping letter-code rows (A, B, C…).
    Returns list of (col_index, name, sublabel).
    """
    profile   = _profile(profile_name)
    xl        = pd.ExcelFile(io.BytesIO(file_bytes))
    start     = 1 if profile.get("skip_sheet_0") else 0
    s_row_idx = profile.get("sublabel_row")   # only GBI has this

    for i in range(start, min(start + 10, len(xl.sheet_names))):
        try:
            df  = xl.parse(i, header=None, na_values=[''])
            raw = df.values.tolist()

            # KP Omni: anchor on Base Weighted, go up 2
            if profile.get("dynamic_base"):
                bw_row = _find_base_weighted_row(raw)
                h_row  = (bw_row - 2) if bw_row is not None else _find_header_row(raw)
            else:
                # All other formats: find the row whose values include 'Total'
                h_row = _find_header_row(raw)
                # Fall back to profile value if Total not found
                if h_row is None:
                    h_row = profile.get("header_row")

            if h_row is None or h_row >= len(raw):
                continue

            hrow = raw[h_row]
            srow = (raw[s_row_idx]
                    if s_row_idx is not None and s_row_idx < len(raw)
                    else [])

            # Determine which column Total lives in → that's col_start
            col_start = 1  # default
            for ci, v in enumerate(hrow):
                if isinstance(v, str) and v.strip().lower() == 'total':
                    col_start = ci
                    break

            cols = []
            for j in range(col_start, len(hrow)):
                v = hrow[j]
                if not isinstance(v, str) or not v.strip() or v.strip() in ('\xa0', ' '):
                    continue
                vs = v.strip()
                # Skip single-letter sig codes
                if len(vs) == 1 and vs.isalpha():
                    continue
                s = (srow[j].strip()
                     if srow and j < len(srow) and isinstance(srow[j], str)
                     else '')
                cols.append((j, vs, s))

            if cols:
                return cols
        except Exception:
            continue
    return []


# ── Sheet parser ──────────────────────────────────────────────


def _find_base_weighted_row(raw):
    """Find row index of 'Base Weighted' in col 0. Returns None if not found."""
    for ri, row in enumerate(raw):
        if row and isinstance(row[0], str):
            cl = row[0].strip().lower()
            if cl in ('base weighted', 'weighted base'):
                return ri
    return None

def _get_kp_omni_statement(raw):
    """For KP Omni, get statement from row 5 if present."""
    if len(raw) > 5 and raw[5] and raw[5][0]:
        v = str(raw[5][0]).strip()
        skip = ['base', 'upper case', 'lower case', 'field dates', 'ipsos', 'table']
        if len(v) > 3 and not any(v.lower().startswith(s) for s in skip):
            return v
    return None

def parse_sheet(file_bytes, sheet_idx, profile_name, col_indices):
    """
    Parse one sheet. Returns wording, base values, answers, values.
    Handles all formats including KP Omni dynamic row detection.
    """
    profile = _profile(profile_name)
    xl      = pd.ExcelFile(io.BytesIO(file_bytes))
    df      = xl.parse(sheet_idx, header=None, na_values=[''])
    raw     = df.values.tolist()

    # Question wording
    q_row   = profile["question_row"]
    wording = str(raw[q_row][0]).strip() if len(raw) > q_row and raw[q_row] and raw[q_row][0] else ''

    # KP: multi-line question wording
    if profile_name == "KP":
        for extra_row in range(q_row + 1, min(q_row + 7, len(raw))):
            next_cell = str(raw[extra_row][0]).strip() if raw[extra_row] and raw[extra_row][0] else ''
            if not next_cell or 'base' in next_cell.lower() or '=' in next_cell:
                break
            wording += ' ' + next_cell

    # KP Omni: grab statement from row 5
    statement = None
    if profile_name == "KP Omni":
        statement = _get_kp_omni_statement(raw)

    # Dynamic base detection (KP Omni)
    if profile.get("dynamic_base"):
        bw_row = _find_base_weighted_row(raw)
        if bw_row is not None:
            b_row      = bw_row
            data_start = bw_row + 1
        else:
            b_row      = 15
            data_start = 16
    else:
        b_row = profile.get("base_row", 7)
        if profile.get("data_start") is not None:
            data_start = profile["data_start"]
        else:
            offset       = profile.get("data_start_base_offset", 2)
            base_row_idx = None
            for ri, row in enumerate(raw):
                if row and isinstance(row[0], str):
                    cl = row[0].strip().lower()
                    if cl.startswith('base') or cl.startswith('unweighted') or 'base =' in cl:
                        base_row_idx = ri
                        break
            data_start = (base_row_idx + offset) if base_row_idx is not None else b_row + 2

    # Base values
    base_data = raw[b_row] if len(raw) > b_row else []
    base_vals = [_coerce(base_data[j] if j < len(base_data) else None) for j in col_indices]

    step    = profile.get("data_step", 3)
    val_off = profile.get("value_row_offset", 1)
    stop_on = set(s.lower() for s in profile.get("stop_on", ["sigma"]))

    main_answers,  main_values,  main_sig  = [], [], []
    net_answers,   net_values,   net_sig   = [], [], []

    i = data_start
    while i < len(raw):
        lbl = raw[i][0] if raw[i] else None
        if isinstance(lbl, str) and lbl.strip():
            cl = lbl.strip().lower()
            if any(s in cl for s in stop_on): break
            if any(cl.startswith(s) for s in ('base:', 'unweighted base', 'base: at least',
                                               'base =', 'weighted base', 'base unweighted',
                                               'upper case', 'lower case', 'field dates')):
                i += 1; continue

            val_row = raw[i + val_off] if i + val_off < len(raw) else []
            sig_row = raw[i + 2]       if i + 2       < len(raw) else []
            vals    = [_coerce(val_row[j] if j < len(val_row) else None) for j in col_indices]
            sigs    = [sig_row[j] if j < len(sig_row) else None for j in col_indices]

            # T2B/B2B rows stay in place; (Net) rows move to bottom
            is_t2b_b2b = ('top 2 box' in cl or 'bottom 2 box' in cl
                          or cl.startswith('t2b') or cl.startswith('b2b'))
            if '(net)' in cl and not is_t2b_b2b:
                net_answers.append(lbl.strip())
                net_values.append(vals)
                net_sig.append(sigs)
            else:
                main_answers.append(lbl.strip())
                main_values.append(vals)
                main_sig.append(sigs)

            i += step
            continue
        i += 1

    # Merge: main rows first, then net rows with has_nets flag so caller can draw separator
    answers  = main_answers + net_answers
    values   = main_values  + net_values
    sig_data = main_sig     + net_sig

    return {
        'wording':   wording,
        'statement': statement,
        'base_vals': base_vals,
        'answers':   answers,
        'values':    values,
        'sig_data':  sig_data,
        'net_start': len(main_answers),   # index where Net rows begin (for separator line)
    }

# ── Excel styles ──────────────────────────────────────────────

HDR_FILL    = PatternFill("solid", fgColor="0F2D4A")
WAVE_FILLS  = [
    PatternFill("solid", fgColor="E8F0FE"),
    PatternFill("solid", fgColor="D2E3FC"),
    PatternFill("solid", fgColor="BCCFEF"),
    PatternFill("solid", fgColor="A8C7FA"),
]
HDR_FONT    = XLFont(bold=True, color="FFFFFF", name="Arial", size=10)
WAVE_FONTS  = [XLFont(bold=True, color="0F2D4A", name="Arial", size=10)] * 4
BODY_FONT   = XLFont(name="Arial", size=10)
CTR         = Alignment(horizontal="center", vertical="center", wrap_text=True)
LEFT        = Alignment(horizontal="left",   vertical="center", wrap_text=True)
THIN        = Side(style="thin",  color="D0D7DE")
BORDER      = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
ALT_FILL    = PatternFill("solid", fgColor="F6F8FA")


def _fmt_pct(v):
    if v is None: return '—'
    if isinstance(v, float):
        pct = math.floor(v*100) if v*100 - math.floor(v*100) < 0.5 else math.ceil(v*100)
        return f"{pct}%"
    return str(v)


NET_BORDER  = Border(left=THIN, right=THIN,
                     top=Side(style="thin", color="0F2D4A"), bottom=THIN)


def _write_table(ws, row, wording, col_headers, base_vals,
                 answers, values, wave_idx=None, net_start=None):
    n_cols = len(col_headers)

    # Title
    c = ws.cell(row=row, column=1, value=wording)
    c.font      = XLFont(bold=True, name="Arial", size=11, color="0F2D4A")
    c.alignment = LEFT
    ws.merge_cells(start_row=row, start_column=1,
                   end_row=row, end_column=n_cols + 1)
    row += 1

    # Header
    use_fill = WAVE_FILLS[wave_idx % len(WAVE_FILLS)] if wave_idx is not None else HDR_FILL
    use_font = WAVE_FONTS[wave_idx % len(WAVE_FONTS)] if wave_idx is not None else HDR_FONT

    ws.cell(row=row, column=1).fill   = use_fill
    ws.cell(row=row, column=1).border = BORDER
    for ci, name in enumerate(col_headers):
        bv   = base_vals[ci] if ci < len(base_vals) else None
        bstr = f"\n(N={int(bv):,})" if isinstance(bv, float) and bv else ''
        cell           = ws.cell(row=row, column=ci + 2, value=name + bstr)
        cell.font      = use_font
        cell.fill      = use_fill
        cell.alignment = CTR
        cell.border    = BORDER
    row += 1

    # Data rows — draw a 1pt dark top border before the first Net row
    has_net_sep = net_start is not None and 0 < net_start < len(answers)
    for ri, answer in enumerate(answers):
        fill       = ALT_FILL if ri % 2 == 0 else PatternFill()
        use_border = NET_BORDER if (has_net_sep and ri == net_start) else BORDER
        lc = ws.cell(row=row, column=1, value=str(answer))
        lc.font=BODY_FONT; lc.alignment=LEFT; lc.border=use_border; lc.fill=fill
        row_vals = values[ri] if ri < len(values) else []
        for ci in range(n_cols):
            v    = row_vals[ci] if ci < len(row_vals) else None
            cell = ws.cell(row=row, column=ci + 2)
            cell.font=BODY_FONT; cell.alignment=CTR
            cell.border=use_border; cell.fill=fill
            cell.value = _fmt_pct(v)
        row += 1

    ws.column_dimensions['A'].width = max(ws.column_dimensions['A'].width, 42)
    for ci in range(n_cols):
        cl = get_column_letter(ci + 2)
        ws.column_dimensions[cl].width = max(ws.column_dimensions[cl].width, 14)

    return row + 2


def _build_toc(wb, toc_entries):
    toc = wb.create_sheet('Contents', 0)
    toc.sheet_view.showGridLines = False
    toc.column_dimensions['A'].width = 6
    toc.column_dimensions['B'].width = 72
    toc.column_dimensions['C'].width = 22

    hf   = PatternFill('solid', fgColor='0F2D4A')
    hfnt = XLFont(bold=True, color='FFFFFF', name='Arial', size=11)
    thin = Side(style='thin', color='D0D7DE')
    brd  = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, val in [(1,'#'), (2,'Question'), (3,'Sheet')]:
        c = toc.cell(row=1, column=col, value=val)
        c.font=hfnt; c.fill=hf; c.border=brd
        c.alignment=Alignment(horizontal='left', vertical='center')
    toc.row_dimensions[1].height = 24

    lf   = XLFont(color='1A6EBD', underline='single', name='Arial', size=10)
    pf   = XLFont(name='Arial', size=10)
    altf = PatternFill('solid', fgColor='F0F5FA')
    et   = Side(style='thin', color='E8ECF0')
    ebrd = Border(left=et, right=et, top=et, bottom=et)

    seen = set()
    r = 2
    for i, (wording, sheet_name) in enumerate(toc_entries):
        if sheet_name in seen: continue
        seen.add(sheet_name)
        fill = altf if i % 2 == 0 else PatternFill()
        safe = sheet_name.replace("'","''")

        nc = toc.cell(row=r, column=1, value=i+1)
        nc.font=pf; nc.border=ebrd; nc.fill=fill
        nc.alignment=Alignment(horizontal='center', vertical='center')

        qc = toc.cell(row=r, column=2, value=wording[:100])
        qc.hyperlink=f"#{safe}!A1"; qc.font=lf
        qc.border=ebrd; qc.fill=fill
        qc.alignment=Alignment(horizontal='left', vertical='center', wrap_text=True)

        sc = toc.cell(row=r, column=3, value=sheet_name)
        sc.font=pf; sc.border=ebrd; sc.fill=fill
        sc.alignment=Alignment(horizontal='left', vertical='center')

        toc.row_dimensions[r].height = 18
        r += 1

    bf = XLFont(color='1A6EBD', underline='single', name='Arial', size=9)
    for ws in wb.worksheets:
        if ws.title == 'Contents': continue
        ws.insert_rows(1)
        nav = ws.cell(row=1, column=1, value='← Contents')
        nav.hyperlink='#Contents!A1'; nav.font=bf
        nav.alignment=Alignment(horizontal='left', vertical='center')


_SUMMARY_MARKERS = frozenset([
    'top 2 box summary', 'top box summary',
    'bottom 2 box summary', 'bottom box summary',
    'grid table', 'summary grid',
    'increased summary', 'decreased summary',
    'mean summary',
])

def _sheet_type_from_text(text):
    tl = text.lower()
    if 'top 2 box' in tl or 'top box summary' in tl:   return 't2b'
    if 'bottom 2 box' in tl or 'bottom box summary' in tl: return 'b2b'
    if 'grid table' in tl or 'summary grid' in tl:      return 'summary_grid'
    if 'mean' in tl or 'average' in tl:                 return 'mean'
    return 'standard'


def read_toc(file_bytes, profile_name):
    """
    Read the TOC/Index sheet (sheet 0) and return a structured question map.

    Returns list of:
        {sheet_idx, prefix, wording, statement, sheet_type, base_text}

    TOC layouts by format
    ─────────────────────
    Corporate Reputation  col0=sheet#  col1=full_text
    GBI / KP Omni / IData col0=sheet#  col1=table#  col2=label  col3=full_text  col4=base
    """
    try:
        df  = pd.ExcelFile(io.BytesIO(file_bytes)).parse(0, header=None, na_values=[''])
        raw = df.values.tolist()
    except Exception:
        return []

    results = []
    corp_rep = (profile_name == "Corporate Reputation")

    for row in raw:
        if not row or row[0] is None:
            continue
        # Col 0 must be a numeric sheet index
        v0 = row[0]
        if isinstance(v0, str):
            try:   v0 = float(v0.strip())
            except: continue
        if not isinstance(v0, (int, float)) or isinstance(v0, bool):
            continue
        if isinstance(v0, float) and math.isnan(v0):
            continue
        sheet_idx = int(v0)

        def cell(ci):
            try:
                v = row[ci]
                return str(v).strip() if v is not None and str(v) not in ('nan', 'None') else ''
            except:
                return ''

        c1, c2, c3, c4 = cell(1), cell(2), cell(3), cell(4)

        if corp_rep:
            full_text = c1
            label     = None
            base_text = ''
        else:
            # GBI / KP Omni / IData:  col2=label, col3=full_text, col4=base
            full_text = c3 if c3 else c1
            label     = c2 if c2 else None
            base_text = c4

        if not full_text or len(full_text) < 3:
            continue

        sheet_type = _sheet_type_from_text(full_text)

        # Split off statement (text after last " - " that isn't a summary marker)
        wording   = full_text
        statement = None
        if ' - ' in full_text:
            head, tail = full_text.rsplit(' - ', 1)
            if not any(m in tail.lower() for m in _SUMMARY_MARKERS):
                wording   = head.strip()
                statement = tail.strip()

        # Extract prefix
        if label and re.match(r'^[A-Za-z0-9_\.]+$', label):
            prefix = label.rstrip('.')
        else:
            m      = re.match(r'^([A-Za-z0-9_]+)[\.\s]', wording)
            prefix = m.group(1) if m else wording[:8]

        results.append({
            'sheet_idx':  sheet_idx,
            'prefix':     prefix,
            'wording':    wording,
            'statement':  statement,
            'sheet_type': sheet_type,
            'base_text':  base_text,
        })

    return results


def toc_to_groups(toc_entries):
    """
    Convert read_toc() list into the same group format as fast_scan():
        {prefix, wording, sheets: [int], types: [str]}
    """
    groups = {}
    order  = []
    for entry in toc_entries:
        p = entry['prefix']
        if p not in groups:
            groups[p] = {
                'prefix':  p,
                'wording': entry['wording'][:100],
                'sheets':  [],
                'types':   set(),
            }
            order.append(p)
        groups[p]['sheets'].append(entry['sheet_idx'])
        groups[p]['types'].add(entry['sheet_type'])

    for g in groups.values():
        g['types'] = sorted(g['types'])
    return [groups[p] for p in order]


def _find_and_parse(file_bytes, prefix, profile_name, col_indices):
    """Find sheets matching prefix (via TOC then fast_scan fallback) and parse the first."""
    toc = read_toc(file_bytes, profile_name)
    if toc:
        matches = [e for e in toc if e['prefix'] == prefix]
        if matches:
            return parse_sheet(file_bytes, matches[0]['sheet_idx'], profile_name, col_indices)
    groups = fast_scan(file_bytes, profile_name)
    match  = next((g for g in groups if g['prefix'] == prefix), None)
    if match is None: return None
    return parse_sheet(file_bytes, match['sheets'][0], profile_name, col_indices)


def _find_and_parse_all(file_bytes, prefix, profile_name, col_indices, include_types=None):
    """
    Like _find_and_parse but returns ALL matching sheets for a prefix,
    filtered by include_types (set of sheet type strings, or None = all).
    Returns list of parsed dicts.
    """
    toc = read_toc(file_bytes, profile_name)
    if toc:
        entries = [e for e in toc if e['prefix'] == prefix]
        if include_types:
            entries = [e for e in entries if e['sheet_type'] in include_types]
        return [p for e in entries
                for p in [parse_sheet(file_bytes, e['sheet_idx'], profile_name, col_indices)]
                if p and p['answers']]

    # Fallback: fast_scan
    groups = fast_scan(file_bytes, profile_name)
    match  = next((g for g in groups if g['prefix'] == prefix), None)
    if match is None: return []
    sheets = match['sheets']
    if include_types:
        types = match.get('types', [])
        sheets = [match['sheets'][i] for i, t in enumerate(types)
                  if t in include_types] if len(types) == len(sheets) else sheets
    return [p for si in sheets
            for p in [parse_sheet(file_bytes, si, profile_name, col_indices)]
            if p and p['answers']]


def generate_excel(selections, files, profile_name, col_indices, col_names,
                   include_types=None):
    """
    include_types: set of sheet type strings to include, e.g. {'standard', 't2b'}.
                   None means include all.
    """
    wb          = openpyxl.Workbook()
    wb.remove(wb.active)
    toc_entries = []

    for sel in selections:
        prefix   = sel['prefix']
        wording  = sel['wording']
        sheet_nm = re.sub(r'[\\/*?\[\]:]', '', prefix)[:31]

        if sheet_nm not in wb.sheetnames:
            wb.create_sheet(title=sheet_nm)
            toc_entries.append((wording[:80], sheet_nm))

        ws  = wb[sheet_nm]
        row = 1 if ws.max_row <= 1 else ws.max_row + 2

        if len(files) == 1:
            parsed_list = _find_and_parse_all(
                files[0]['bytes'], prefix, profile_name, col_indices, include_types)
            for p in parsed_list:
                row = _write_table(ws, row, p['wording'], col_names,
                                   p['base_vals'], p['answers'], p['values'],
                                   net_start=p.get('net_start'))
        else:
            title_cell = ws.cell(row=row, column=1, value=wording[:100])
            title_cell.font      = XLFont(bold=True, name="Arial", size=12, color="0F2D4A")
            title_cell.alignment = LEFT
            ws.merge_cells(start_row=row, start_column=1,
                           end_row=row, end_column=len(col_names) + 1)
            row += 2

            for fi, finfo in enumerate(files):
                parsed_list = _find_and_parse_all(
                    finfo['bytes'], prefix, profile_name, col_indices, include_types)
                if parsed_list:
                    for p in parsed_list:
                        row = _write_table(ws, row, finfo['label'], col_names,
                                           p['base_vals'], p['answers'], p['values'],
                                           wave_idx=fi, net_start=p.get('net_start'))
                else:
                    # Question not found in this wave — write a placeholder row
                    na_fill = WAVE_FILLS[fi % len(WAVE_FILLS)]
                    na_font = WAVE_FONTS[fi % len(WAVE_FONTS)]
                    lbl = ws.cell(row=row, column=1,
                                  value=f"{finfo['label']}  —  not available in this file")
                    lbl.font      = na_font
                    lbl.fill      = na_fill
                    lbl.alignment = LEFT
                    lbl.border    = BORDER
                    ws.merge_cells(start_row=row, start_column=1,
                                   end_row=row, end_column=len(col_names) + 1)
                    row += 3

    if toc_entries:
        _build_toc(wb, toc_entries)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Word / Media Release export ───────────────────────────────

# Resolve template path relative to this file so it works on Streamlit Cloud
TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'template_doc.docx')
BRAND_COLOR   = RGBColor(0x2F, 0x46, 0x9C)


def _add_run(para, text, bold=False, size_pt=None, color=None):
    run = para.add_run(text)
    run.bold = bold
    if size_pt:  run.font.size = Pt(size_pt)
    if color:    run.font.color.rgb = color
    return run


def _insert_para_after(ref_para, text='', bold=False, size_pt=None, color=None):
    new_p = OxmlElement('w:p')
    ref_para._element.addnext(new_p)
    fake = DocxPara(new_p, ref_para._parent)
    if text:
        _add_run(fake, text, bold=bold, size_pt=size_pt, color=color or BRAND_COLOR)
    return fake


def _write_word_table(doc, insert_after_para, question_wording,
                      col_names, base_vals, answers, values):

    q_para = _insert_para_after(insert_after_para, question_wording,
                                 bold=True, size_pt=10, color=BRAND_COLOR)
    spacer = _insert_para_after(q_para)

    n_cols = len(col_names)
    tbl_el = OxmlElement('w:tbl')
    spacer._element.addnext(tbl_el)

    tblPr = OxmlElement('w:tblPr')
    tblW  = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0'); tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    tbl_el.append(tblPr)

    def make_cell(text, bold=False, center=True, bg_hex=None, size_pt=9):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        if bg_hex:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_hex)
            tcPr.append(shd)
        tc.append(tcPr)
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        if center:
            jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
        sp = OxmlElement('w:spacing'); sp.set(qn('w:after'), '0'); pPr.append(sp)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b'); rPr.append(b)
        col_el = OxmlElement('w:color')
        col_el.set(qn('w:val'), 'FFFFFF' if bg_hex == '2F469C' else '000000')
        rPr.append(col_el)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size_pt * 2)); rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = str(text) if text else ''
        if text and (str(text).startswith(' ') or str(text).endswith(' ')):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t); p.append(r); tc.append(p)
        return tc

    hdr_tr = OxmlElement('w:tr')
    hdr_tr.append(make_cell('', bold=True, center=True, bg_hex='2F469C'))
    for ci, name in enumerate(col_names):
        bv   = base_vals[ci] if ci < len(base_vals) else None
        bstr = f'\n(N={int(bv):,})' if isinstance(bv, float) and bv else ''
        hdr_tr.append(make_cell(name + bstr, bold=True, center=True, bg_hex='2F469C'))
    tbl_el.append(hdr_tr)

    for ri, answer in enumerate(answers):
        data_tr = OxmlElement('w:tr')
        bg = 'F6F8FA' if ri % 2 == 0 else None
        data_tr.append(make_cell(answer, bold=False, center=False, bg_hex=bg))
        row_vals = values[ri] if ri < len(values) else []
        for ci in range(n_cols):
            v = row_vals[ci] if ci < len(row_vals) else None
            data_tr.append(make_cell(_fmt_pct(v), center=True, bg_hex=bg))
        tbl_el.append(data_tr)

    return tbl_el


def generate_word(selections, files, profile_name, col_indices, col_names,
                  survey_title='', include_types=None):
    import os
    if not os.path.exists(TEMPLATE_PATH):
        return None, "template_doc.docx not found in app directory"

    doc = Document(TEMPLATE_PATH)

    title_para = doc.paragraphs[1]
    for run in title_para.runs:
        run.text = ''
    if title_para.runs:
        title_para.runs[0].text = survey_title or 'Topline Findings'
    else:
        _add_run(title_para, survey_title or 'Topline Findings',
                 bold=True, color=BRAND_COLOR)

    topline_para = next(
        (p for p in doc.paragraphs if 'Topline' in p.text), doc.paragraphs[-1]
    )
    insert_after = topline_para

    for sel in selections:
        prefix = sel['prefix']
        for fi, finfo in enumerate(files):
            parsed_list = _find_and_parse_all(
                finfo['bytes'], prefix, profile_name, col_indices, include_types)
            if parsed_list:
                for p in parsed_list:
                    title  = (f"{p['wording']}  —  {finfo['label']}"
                              if len(files) > 1 else p['wording'])
                    tbl_el = _write_word_table(doc, insert_after, title,
                                               col_names, p['base_vals'],
                                               p['answers'], p['values'])
                    spacer_el = OxmlElement('w:p')
                    tbl_el.addnext(spacer_el)
                    insert_after = DocxPara(spacer_el, topline_para._parent)
            elif len(files) > 1:
                # Placeholder paragraph for missing wave
                na_para = _insert_para_after(
                    insert_after,
                    f"{finfo['label']}  —  not available in this file",
                    bold=True, size_pt=9,
                )
                insert_after = na_para

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), None


def deep_scan_file(file_bytes):
    """
    Full-file structural scan for unknown banner formats.

    Reads the TOC (if present), samples up to 8 sheets spread across the
    whole file, infers the dominant row layout, and flags outlier sheets.

    Returns dict:
        has_toc        bool
        toc_summary    str
        sheet_count    int
        sampled        list of per-sheet dicts
        dominant       dict with question_row/header_row/base_row/data_start/
                           data_step/col_start/stop_on  (all 0-indexed)
        outlier_sheets list of sheet indices that don't match dominant
        variant_count  int  (number of distinct non-dominant patterns)
        row_preview    list[list[str]]  first 15 rows × 12 cols of first data sheet
    """
    try:
        xl          = pd.ExcelFile(io.BytesIO(file_bytes))
        sheet_names = xl.sheet_names
        n_sheets    = len(sheet_names)
    except Exception as e:
        return {'error': str(e), 'sheet_count': 0}

    # ── TOC detection ─────────────────────────────────────────
    has_toc     = False
    toc_summary = ''
    if n_sheets > 1:
        try:
            df0  = xl.parse(0, header=None, nrows=6, na_values=[''])
            raw0 = df0.values.tolist()
            toc_kw = ('sheet number', 'table of contents', 'contents', 'index', 'table number')
            for row in raw0[:4]:
                if row and row[0] and isinstance(row[0], str):
                    if any(kw in row[0].strip().lower() for kw in toc_kw):
                        has_toc = True
                        break
            if not has_toc:
                num_count = sum(
                    1 for row in raw0[1:]
                    if row and row[0] is not None
                    and isinstance(row[0], (int, float))
                    and not isinstance(row[0], bool)
                )
                has_toc = num_count >= 2
            if has_toc:
                df0f        = xl.parse(0, header=None, na_values=[''])
                entry_count = sum(
                    1 for row in df0f.values.tolist()
                    if row and isinstance(row[0], (int, float))
                    and not isinstance(row[0], bool)
                )
                toc_summary = f"TOC found — ~{entry_count} entries"
        except Exception:
            pass

    # ── Sheet sampling ────────────────────────────────────────
    start       = 1 if (has_toc and n_sheets > 1) else 0
    data_sheets = list(range(start, n_sheets))
    if not data_sheets:
        data_sheets = [0]

    # Always include the first 15 data sheets for reliable detection,
    # then append a few spread across the rest to catch outliers.
    front          = data_sheets[:15]
    if len(data_sheets) > 15:
        tail_step  = max(1, (len(data_sheets) - 15) // 5)
        tail       = data_sheets[15::tail_step][:5]
    else:
        tail       = []
    sample_indices = sorted(set(front + tail))

    sampled        = []
    sheet_previews = []   # up to 3 sheets shown in the wizard UI

    base_kws = ('base', 'unweighted base', 'base weighted', 'weighted base', 'n =', 'n=')

    for si in sample_indices:
        try:
            df  = xl.parse(si, header=None, nrows=40, na_values=[''])
            raw = df.values.tolist()
            if not raw or len(raw) < 4:
                continue

            cleaned = [
                [(str(c) if c is not None and str(c) not in ('nan', 'None') else '')
                 for c in (row[:10] if len(row) >= 10 else row + [''] * (10 - len(row)))]
                for row in raw[:18]
            ]
            if len(sheet_previews) < 3:
                sheet_previews.append({
                    'sheet_idx':  si,
                    'sheet_name': sheet_names[si],
                    'rows':       cleaned,
                })

            h_row = _find_header_row(raw, max_rows=25)

            # Base row
            b_row = None
            for ri, row in enumerate(raw):
                if row and isinstance(row[0], str):
                    cl = row[0].strip().lower()
                    if any(cl == kw or cl.startswith(kw + ' ') or cl.startswith(kw + ':')
                           for kw in base_kws):
                        b_row = ri
                        break

            # Question row — last long string above h_row
            q_row = None
            if h_row is not None:
                for candidate in range(h_row - 1, max(-1, h_row - 6), -1):
                    if 0 <= candidate < len(raw):
                        v = raw[candidate][0] if raw[candidate] else None
                        if v and isinstance(v, str):
                            vs = v.strip()
                            if (len(vs) > 10
                                    and not vs.lower().startswith('table ')
                                    and not vs.replace('.', '').replace(',', '').replace(' ', '').isnumeric()):
                                q_row = candidate
                                break

            # Data start
            data_start = None
            if b_row is not None:
                data_start = b_row + 1
                if data_start < len(raw) and raw[data_start]:
                    v = raw[data_start][0]
                    # If the row right after base is very short, skip it (sublabel row)
                    if v and isinstance(v, str) and len(v.strip()) < 4:
                        data_start = b_row + 2

            # Column start (where "Total" lives)
            col_start = 1
            if h_row is not None and h_row < len(raw):
                for ci, v in enumerate(raw[h_row]):
                    if isinstance(v, str) and v.strip().lower() == 'total':
                        col_start = ci
                        break

            # Stop-row patterns found in this sheet
            stop_patterns = []
            known_stops   = ('sigma', 'back to top', 'total mentions',
                             'overlap formula used', 'field dates')
            for row in raw:
                if row and row[0] and isinstance(row[0], str):
                    v = row[0].strip().lower()
                    if v in known_stops and v not in stop_patterns:
                        stop_patterns.append(v)

            sampled.append({
                'sheet_idx':     si,
                'sheet_name':    sheet_names[si],
                'question_row':  q_row,
                'header_row':    h_row,
                'base_row':      b_row,
                'data_start':    data_start,
                'data_step':     3,
                'col_start':     col_start,
                'stop_patterns': stop_patterns,
            })
        except Exception:
            continue

    if not sampled:
        return {
            'has_toc': has_toc, 'toc_summary': toc_summary,
            'sheet_count': n_sheets, 'sampled': [],
            'dominant': None, 'outlier_sheets': [], 'variant_count': 0,
            'sheet_previews': sheet_previews,
        }

    # ── Find dominant pattern ─────────────────────────────────
    def sig(s):
        return (s['question_row'], s['header_row'], s['base_row'],
                s['data_start'], s['col_start'])

    valid      = [s for s in sampled
                  if all(v is not None for v in [s['question_row'], s['header_row'],
                                                 s['base_row'], s['data_start']])]
    sig_counts = Counter(sig(s) for s in valid)

    if sig_counts:
        dom_sig, _  = sig_counts.most_common(1)[0]
        all_stops   = set()
        for s in valid:
            if sig(s) == dom_sig:
                all_stops.update(s['stop_patterns'])
        dominant = {
            'question_row': dom_sig[0],
            'header_row':   dom_sig[1],
            'base_row':     dom_sig[2],
            'data_start':   dom_sig[3],
            'data_step':    3,
            'col_start':    dom_sig[4],
            'stop_on':      sorted(all_stops) or ['sigma'],
        }
        outlier_sheets = [s['sheet_idx'] for s in valid if sig(s) != dom_sig]
        variant_count  = max(0, len(sig_counts) - 1)
    else:
        s0       = sampled[0]
        dominant = {
            'question_row': s0.get('question_row') or 2,
            'header_row':   s0.get('header_row') or 4,
            'base_row':     s0.get('base_row') or 7,
            'data_start':   s0.get('data_start') or 9,
            'data_step':    3,
            'col_start':    s0.get('col_start') or 1,
            'stop_on':      s0.get('stop_patterns') or ['sigma'],
        }
        outlier_sheets = []
        variant_count  = 0

    return {
        'has_toc':        has_toc,
        'toc_summary':    toc_summary,
        'sheet_count':    n_sheets,
        'sampled':        sampled,
        'dominant':       dominant,
        'outlier_sheets': outlier_sheets,
        'variant_count':  variant_count,
        'sheet_previews': sheet_previews,
    }


def detect_and_describe(file_bytes):
    try:
        xl       = pd.ExcelFile(io.BytesIO(file_bytes))
        n_sheets = len(xl.sheet_names)
    except Exception as e:
        return {'matched_profile': None, 'findings': [('Error', str(e), 'warn')],
                'sample_columns': [], 'n_sheets': 0}

    findings = [('Total sheets', str(n_sheets), 'info')]

    ref_si = None
    for si in range(min(6, n_sheets)):
        try:
            df  = xl.parse(si, header=None, nrows=20, na_values=[''])
            raw = df.values.tolist()

            # Skip if this looks like a TOC/index (short strings, no question wording)
            c2 = str(raw[2][0]).strip() if len(raw)>2 and raw[2] and raw[2][0] else ''
            c4 = str(raw[4][0]).strip() if len(raw)>4 and raw[4] and raw[4][0] else ''

            # KP Omni: row 2 = "Table N", row 4 = real question wording
            if c2.lower().startswith('table ') and len(c4) > 10:
                ref_si = si; ref_raw = raw; break

            # Other formats: row 2 has real question wording (long string)
            if len(c2) > 10 and not c2.lower().startswith('sheet') and not c2.lower().startswith('table'):
                ref_si = si; ref_raw = raw; break

        except: continue

    if ref_si is None:
        return {'matched_profile': None,
                'findings': findings + [('Data sheets', 'None found', 'warn')],
                'sample_columns': [], 'n_sheets': n_sheets}

    def cell(r, c):
        try:
            v = ref_raw[r][c]
            return str(v).strip() if v and str(v) not in ('nan','None') else ''
        except: return ''

    def row_has_total(r):
        try:
            return any('total' in str(v).lower()
                       for v in (ref_raw[r] or [])
                       if v and str(v) not in ('nan','None'))
        except: return False

    r0=cell(0,0); r1=cell(1,0); r2=cell(2,0)
    r3c0=cell(3,0); r3c1=cell(3,1)
    r4c0=cell(4,0); r4c1=cell(4,1)
    r5c1=cell(5,1)

    findings.append(('First data sheet', f'Sheet {ref_si} of {n_sheets}', 'info'))

    # KP Omni: row 2 = "Table N", row 4 = question, has Base Weighted
    if r2.lower().startswith('table') and len(r4c0) > 10:
        has_bw = any('base weighted' in str(ref_raw[ri][0] if ref_raw[ri] else '').lower()
                     for ri in range(min(20, len(ref_raw))))
        if has_bw:
            bw_row = _find_base_weighted_row(ref_raw)
            h_row  = (bw_row - 2) if bw_row else None
            cols   = [str(v).strip() for v in (ref_raw[h_row] if h_row else [])
                      if v and str(v) not in ('nan','None','\xa0')]
            findings += [
                ('Project name',   f'Row 1: "{r1[:50]}"', 'ok'),
                ('Table number',   f'Row 2: "{r2}"', 'ok'),
                ('Question wording', f'Row 4: "{r4c0[:50]}"', 'ok'),
                ('Column headers', f'Row {h_row}: {cols[:4]}', 'ok'),
                ('Base Weighted',  f'Row {bw_row}', 'ok'),
                ('Data start',     f'Row {bw_row+1 if bw_row else "?"} (every 3 rows)', 'ok'),
            ]
            return {'matched_profile': 'KP Omni',
                    'findings': findings, 'sample_columns': cols, 'n_sheets': n_sheets}

    # Corporate Reputation
    if (len(r3c0) > 10
            and ('total sample' in r2.lower() or 'weight' in r2.lower() or len(r2) < 50)
            and ('total' in r4c1.lower() or 'total' in r5c1.lower())):
        h_row  = 5 if 'total' in r5c1.lower() else 4
        h_data = ref_raw[h_row] if len(ref_raw) > h_row else []
        cols   = [str(v).strip() for v in h_data[1:] if v and str(v) not in ('nan','None','\xa0')]
        findings += [
            ('Question wording',   f'Row 3: "{r3c0[:50]}"', 'ok'),
            ('Descriptor row',     f'Row 2: "{r2[:40]}"', 'ok'),
            ('Column headers row', f'Row {h_row}: {cols[:4]}', 'ok'),
            ('Base row',           'Row 8 (Unweighted Base)', 'ok'),
            ('Data start',         'Row 12 (every 3 rows)', 'ok'),
        ]
        return {'matched_profile': 'Corporate Reputation',
                'findings': findings, 'sample_columns': cols, 'n_sheets': n_sheets}

    # IData: row 2 = question, row 3 col 0 blank (group headers), row 4 col 1 = Total
    # Distinguished from GBI by Total NOT appearing in row 3
    if (len(r2) > 10
            and not r3c0
            and 'total' not in (r3c0 + r3c1).lower()
            and 'total' in r4c1.lower()):
        h_data = ref_raw[4] if len(ref_raw) > 4 else []
        cols   = [str(v).strip() for v in h_data[1:]
                  if v and str(v) not in ('nan','None','\xa0')]
        findings += [
            ('Question wording', f'Row 2: "{r2[:50]}"', 'ok'),
            ('Group headers',    f'Row 3: (blank col 0)', 'ok'),
            ('Column headers',   f'Row 4: {cols[:4]}', 'ok'),
            ('Base row',         'Row 7 (Base: Total Answering)', 'ok'),
            ('Data start',       'Row 9 (every 3 rows)', 'ok'),
        ]
        return {'matched_profile': 'IData',
                'findings': findings, 'sample_columns': cols, 'n_sheets': n_sheets}

    # Global Brand Identity
    if (len(r2) > 10
            and ('total' in r3c0.lower() or 'total' in r3c1.lower())
            and 'total' in r4c1.lower()):
        col_start = 0 if 'total' in r3c0.lower() else 1
        h_data = ref_raw[3] if len(ref_raw) > 3 else []
        cols   = [str(v).strip() for v in h_data[col_start:]
                  if v and str(v) not in ('nan','None','\xa0')]
        findings += [
            ('Question wording', f'Row 2: "{r2[:50]}"', 'ok'),
            ('Country headers',  f'Row 3: {cols[:5]}', 'ok'),
            ('Base row',         'Row 7 (Unweighted Base)', 'ok'),
            ('Data start',       'Base row + 2 (every 3 rows)', 'ok'),
        ]
        return {'matched_profile': 'Global Brand Identity',
                'findings': findings, 'sample_columns': cols, 'n_sheets': n_sheets}

    # KP
    if len(r2) > 10 and row_has_total(1):
        h_data = ref_raw[4] if len(ref_raw) > 4 else []
        cols   = [str(v).strip() for v in h_data[1:] if v and str(v) not in ('nan','None','\xa0')]
        findings += [
            ('Question wording', f'Row 2: "{r2[:50]}"', 'ok'),
            ('Column headers',   f'Row 4: {cols[:4]}', 'ok'),
            ('Base row',         'Row 7', 'ok'),
            ('Data start',       'Dynamic (after base row)', 'ok'),
        ]
        return {'matched_profile': 'KP',
                'findings': findings, 'sample_columns': cols, 'n_sheets': n_sheets}

    findings += [
        ('Row 2', r2[:50] or '(empty)', 'info'),
        ('Row 3 col 0', r3c0[:40] or '(empty)', 'info'),
        ('Row 4 col 0', r4c0[:40] or '(empty)', 'info'),
        ('Tip', 'Run examine_structure.py in Colab and share the output', 'warn'),
    ]
    return {'matched_profile': None,
            'findings': findings, 'sample_columns': [], 'n_sheets': n_sheets}
